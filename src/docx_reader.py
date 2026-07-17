from docx import Document
from docx.document import Document as DocumentType


def load_document(path: str) -> DocumentType:
    return Document(path)


def iter_paragraphs(doc: DocumentType):
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

